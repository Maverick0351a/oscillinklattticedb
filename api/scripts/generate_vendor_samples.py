"""Generate small vendor sample files and update bench/assets/manifest.yml.

This script creates a tiny corpus under bench/assets/vendor (docx, csv, html, eml, txt),
computes SHA-256 for each, and updates bench/assets/manifest.yml so downstream steps can
fetch/extract uniformly. CSV generation falls back to a deterministic synthetic dataset
if scikit-learn is unavailable.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
from pathlib import Path
from typing import Any, Dict, List

import yaml


def sha256_of(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            h.update(chunk)
    return h.hexdigest()


def gen_docx(dst: Path) -> Dict[str, Any]:
    try:
        from docx import Document  # type: ignore
    except Exception as e:  # pragma: no cover - optional dependency
        return {"ok": False, "error": f"python-docx not available: {e}"}
    doc = Document()
    doc.add_heading("Oscillink Sample Document", level=1)
    doc.add_paragraph(
        "This document was generated to seed the asset corpus for benchmarks."
    )
    doc.add_paragraph("It includes a short paragraph and a list:")
    for item in ["Routing", "Composition", "Receipts", "Determinism"]:
        doc.add_paragraph(item, style="List Bullet")
    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(dst))
    return {"ok": True, "path": str(dst)}


def gen_csv(dst: Path) -> Dict[str, Any]:
    # Try iris from scikit-learn, otherwise synthesize a tiny dataset
    headers: List[str]
    rows: List[List[Any]] = []
    try:  # pragma: no cover - optional dependency
        from sklearn.datasets import load_iris  # type: ignore
        # Prefer explicit tuple return for clear typing; if this fails, we'll fallback entirely
        X, y = load_iris(return_X_y=True)
        feature_names = [
            "sepal length (cm)",
            "sepal width (cm)",
            "petal length (cm)",
            "petal width (cm)",
        ]
        headers = [*feature_names, "target"]
        for row, t in zip(X.tolist(), y.tolist()):
            # round floats to 3 decimals for determinism
            rows.append([*(f"{v:.3f}" for v in row), int(t)])
    except Exception:
        headers = ["f1", "f2", "f3", "f4", "target"]
        # Deterministic small grid
        for i in range(10):
            rows.append([f"{0.1 * i:.3f}", f"{0.2 * i:.3f}", f"{0.3 * i:.3f}", f"{0.4 * i:.3f}", i % 3])

    dst.parent.mkdir(parents=True, exist_ok=True)
    with dst.open("w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(headers)
        w.writerows(rows)
    return {"ok": True, "path": str(dst)}


def gen_html(dst: Path) -> Dict[str, Any]:
    html = (
        "<!doctype html><html><head><meta charset=\"utf-8\"><title>Oscillink Sample</title></head>"
        "<body><h1>Oscillink HTML Sample</h1><p>This is a tiny HTML file for ingestion.</p>"
        "<ul><li>Alpha</li><li>Beta</li></ul></body></html>"
    )
    dst.parent.mkdir(parents=True, exist_ok=True)
    dst.write_text(html, encoding="utf-8")
    return {"ok": True, "path": str(dst)}


def gen_eml(dst: Path) -> Dict[str, Any]:
    eml = (
        "From: sender@example.com\n"
        "To: receiver@example.com\n"
        "Subject: Oscillink Sample Email\n\n"
        "Hello, this is a simple RFC-822 style message body for testing.\n"
    )
    dst.parent.mkdir(parents=True, exist_ok=True)
    dst.write_text(eml, encoding="utf-8")
    return {"ok": True, "path": str(dst)}


def gen_txt(dst: Path) -> Dict[str, Any]:
    text = (
        "Vendor sample TXT file.\n"
        "This text is here to exercise ingestion and chunking paths for Oscillink LatticeDB.\n"
    )
    dst.parent.mkdir(parents=True, exist_ok=True)
    dst.write_text(text, encoding="utf-8")
    return {"ok": True, "path": str(dst)}


def update_manifest(manifest_path: Path, base_dir: Path, entries: List[Dict[str, Any]]) -> None:
    data: Dict[str, Any] = {}
    if manifest_path.exists():
        data = yaml.safe_load(manifest_path.read_text()) or {}
    if not data:
        data = {"base_dir": base_dir.as_posix(), "files": []}
    data.setdefault("base_dir", base_dir.as_posix())
    files: List[Dict[str, Any]] = data.setdefault("files", [])

    # Index by relative path for idempotent updates
    idx = {item.get("path"): i for i, item in enumerate(files) if item.get("path")}
    for e in entries:
        if e["path"] in idx:
            i = idx[e["path"]]
            files[i]["sha256"] = e["sha256"]
            files[i]["url"] = e.get("url")  # may remain None
        else:
            files.append({"url": e.get("url"), "sha256": e["sha256"], "path": e["path"]})

    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    manifest_path.write_text(yaml.safe_dump(data, sort_keys=False))


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--root", default="bench/assets/vendor", help="Output directory for vendor files")
    ap.add_argument("--base-dir", default="bench/assets", help="Base directory whose relative paths are recorded in manifest")
    ap.add_argument("--manifest", default="bench/assets/manifest.yml", help="Manifest YAML path to update")
    args = ap.parse_args()

    base_dir = Path(args.base_dir).resolve()
    root = Path(args.root)
    manifest_path = Path(args.manifest)

    outputs = [
        (root / "docx" / "generated.docx", gen_docx),
        (root / "csv" / "iris.csv", gen_csv),
        (root / "html" / "sample.html", gen_html),
        (root / "eml" / "sample.eml", gen_eml),
        (root / "txt" / "sample.txt", gen_txt),
    ]

    created: List[Dict[str, Any]] = []
    for dst, fn in outputs:
        res = fn(dst)
        if not res.get("ok", False):
            # Skip files we cannot build but keep going
            continue
        h = sha256_of(dst)
        rel = dst.resolve().relative_to(base_dir).as_posix()
        created.append({"path": rel, "sha256": h, "url": None})

    if created:
        update_manifest(manifest_path, base_dir, created)

    print(json.dumps({"written": len(created), "base_dir": base_dir.as_posix(), "items": created}, indent=2))


if __name__ == "__main__":  # pragma: no cover - CLI
    main()
